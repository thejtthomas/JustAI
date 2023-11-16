import streamlit as st
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx2pdf import convert

def generate_word_document(html_content):
    # Create a BytesIO object to store the Word document
    word_output = BytesIO()

    # Create a new Word document
    doc = Document()

    # Add a heading
    doc.add_heading("Generated Word Document", level=1)

    # Parse HTML content and add paragraphs to the Word document
    # This is a simple example, and you might need a more sophisticated HTML-to-docx conversion
    # depending on your HTML structure
    paragraphs = html_content.split('<br>')
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph.strip())

    # Save the Word document to the BytesIO object
    doc.save(word_output)

    # Seek to the beginning of the BytesIO stream
    word_output.seek(0)

    return word_output

# Streamlit UI
st.title("HTML to Word Document Generator")

# Text area for entering HTML content
html_content = st.text_area("Enter HTML Content")

# Button to generate Word document
if st.button("Generate Word Document"):
    # Check if HTML content is provided
    if html_content:
        # Generate Word document
        word_output = generate_word_document(html_content)

        # Provide a download link for the Word document
        st.download_button(
            label="Download Word Document",
            data=word_output,
            file_name="generated_document.docx",
            key="word-doc-download"
        )
    else:
        st.warning("Please enter HTML content before generating the Word document.")
